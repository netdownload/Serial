# TODO Создать отчет в Word


import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Cm
from calendar import monthrange
# Библиотека для работы с датой (для расчета разницы во времени и прибавления времени)
import datetime

import pymysql
import sys
# Библиотека для работы с логаи
import logging
import logging.handlers

SDO_TEXT = '2,1'  # Суточный договорной объем
SDO = 2.1
MSO_TEXT = '2,310'  # Максимальный суточный объем
MSO = 2.310
TEST_DATE_BEGIN = '01.04.2020'
TEST_DATE_END = '01.05.2020'
DATABASE_HOST = '10.1.1.50'
DATABASE_USER = 'user'
DATABASE_PASSWORD = 'qwerty123'
DATABASE = 'resources'


def format_text(cell_t_b, cell_t_e, row_t_b, row_t_e, font_t, bold_t, italic_t, size_t, table_alignment,
                vertical_alignment):
    t_al = WD_TABLE_ALIGNMENT.CENTER
    v_al = WD_ALIGN_VERTICAL.CENTER
    if table_alignment == 'center':
        t_al = WD_TABLE_ALIGNMENT.CENTER
    if vertical_alignment == 'center':
        v_al = WD_ALIGN_VERTICAL.CENTER
    for row_t in range(row_t_b, row_t_e + 1):
        for cell_t in range(cell_t_b, cell_t_e + 1):
            table.rows[row_t].cells[cell_t].paragraphs[0].runs[0].font.bold = bold_t
            table.rows[row_t].cells[cell_t].paragraphs[0].runs[0].font.italic = italic_t
            table.rows[row_t].cells[cell_t].paragraphs[0].runs[0].font.name = font_t
            table.rows[row_t].cells[cell_t].paragraphs[0].runs[0].font.size = Pt(size_t)
            table.cell(row_t, cell_t).paragraphs[0].paragraph_format.alignment = t_al
            table.cell(row_t, cell_t).vertical_alignment = v_al
    return 0


# Функция для проверки доступности базы данных
def check_database_connection():
    try:
        pymysql.connect(host='10.1.1.50',
                        user='user',
                        password='qwerty123',
                        db='resources')
        logging.debug('База данных доступна')
    except pymysql.err.OperationalError:
        logging.debug('База данных недоступна')
        sys.exit("База данных недоступна")


# Функция берет последнее значение времени из базы
def get_gas_value_from_database(date_sql_begin, date_sql_end):
    connection = pymysql.connect(host='10.1.1.50',
                                 user='user',
                                 password='qwerty123',
                                 db='resources')

    delta = datetime.datetime.strptime(date_sql_end, '%d.%m.%Y') - datetime.datetime.strptime(date_sql_begin,
                                                                                              '%d.%m.%Y')
    date_sql = datetime.datetime.strptime(date_sql_begin, '%d.%m.%Y') + datetime.timedelta(hours=10)
    gas_v_st = []
    try:
        with connection.cursor() as cursor:
            for days_i in range(0, delta.days+1):
                print(date_sql)
                sql = 'SELECT gas_v_st_s FROM gas WHERE gas_datetime=%s LIMIT 0, 1'
                cursor.execute(sql, date_sql)
                result = cursor.fetchone()
                gas_v_st.append(result[0])
                date_sql = date_sql + datetime.timedelta(hours=24)
    finally:
        connection.close()
    return gas_v_st


gas_value = get_gas_value_from_database(TEST_DATE_BEGIN, TEST_DATE_END)
print(gas_value[0])

doc = docx.Document()

days_in_month = monthrange(2020, 3)[1]

par = doc.add_paragraph()
par.alignment = WD_ALIGN_PARAGRAPH.CENTER
r1 = par.add_run('Приложение к Акту поданного-принятого газа\n')
font = r1.font
font.size = Pt(12)
font.name = 'Times New Roman'
font.bold = True
r2 = par.add_run('между ООО «Газпром межрегионгаз Иваново» и ОАО «Аньковское»')
font = r2.font
font.name = 'Times New Roman'

par = doc.add_paragraph()
par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r1 = par.add_run('от 1 Мая 2020 г.')  # TODO Дата должна браться исходя из периода отчета
font = r1.font
font.size = Pt(12)
font.name = 'Times New Roman'
font.underline = True

par = doc.add_paragraph()
par.alignment = WD_ALIGN_PARAGRAPH.CENTER
r1 = par.add_run('За Апрель месяц 2020 года')  # TODO Дата должна браться исходя из периода отчета
font = r1.font
font.size = Pt(12)
font.name = 'Times New Roman'

par = doc.add_paragraph()
par.alignment = WD_ALIGN_PARAGRAPH.LEFT
# TODO Количество берется из базы данных
r1 = par.add_run('Принято всего 40,641 тыс. н.м.³, в том числе за каждые сутки месяца:')
font = r1.font
font.size = Pt(12)
font.name = 'Times New Roman'

# def make_rows_bold(*rows):
#     for row in rows:
#         for cell in row.cells:
#             for paragraph in cell.paragraphs:
#                 for run in paragraph.runs:
#                     run.font.bold = True

# make_rows_bold(table.rows[0])

# Настройка таблицы
table = doc.add_table(rows=days_in_month + 3, cols=5)
table.style = 'Table Grid'
table.allow_autofit = False
table.cell(0, 0).width = Cm(3)
table.cell(0, 1).width = Cm(3.8)
table.cell(0, 2).width = Cm(3.8)
table.cell(0, 3).width = Cm(3.8)
table.cell(0, 4).width = Cm(3.8)
# Шапка таблицы
table.cell(0, 0).text = 'Дата'
table.cell(0, 1).text = 'Суточный договорной объем'
table.cell(0, 2).text = 'Максимальный суточный объем'
table.cell(0, 3).text = 'Фактический объем принятого газа'
table.cell(0, 4).text = 'Перерасход газа за каждые сутки от максимального суточного объема'
table.cell(1, 0).text = '1'
table.cell(1, 1).text = '2'
table.cell(1, 2).text = '3'
table.cell(1, 3).text = '4'
table.cell(1, 4).text = '5'
# Форматирование ячеек
format_text(0, 4, 0, 1, 'Times New Roman', True, False, 11, 'center', 'center')
# Вывод дней
for days in range(1, days_in_month + 1):
    table.cell(days + 1, 0).text = str(days)
format_text(0, 0, 2, days_in_month + 1, 'Times New Roman', True, True, 11, 'center', 'center')
# Вывод суточного договорного объема
for days in range(1, days_in_month + 1):
    table.cell(days + 1, 1).text = SDO_TEXT
format_text(1, 1, 2, days_in_month + 1, 'Times New Roman', False, False, 11, 'center', 'center')
# Вывод максимального суточного объема
for days in range(1, days_in_month + 1):
    table.cell(days + 1, 2).text = MSO_TEXT
format_text(2, 2, 2, days_in_month + 1, 'Times New Roman', False, False, 11, 'center', 'center')
doc.save('Prilozhenie.docx')
